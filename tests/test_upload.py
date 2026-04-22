import requests
import json
import os

BASE_URL = "http://localhost:5000/api"
TOKEN = ""


def print_result(test_name, response):
    print(f"\n{'='*50}")
    print(f"🔍 {test_name}")
    print(f"   Status Code : {response.status_code}")
    try:
        print(f"   Response    : {json.dumps(response.json(), indent=2)}")
    except:
        print(f"   Response    : {response.text[:300]}")
    print(f"{'='*50}")


# ─── LOGIN ────────────────────────────────────────────────

def login():
    global TOKEN
    response = requests.post(f"{BASE_URL}/auth/login", json={
        "email": "testuser@studygpt.com",
        "password": "Test1234!"
    })
    if response.status_code == 200:
        TOKEN = response.json().get("token")
        print(f"\n✅ Logged in. Token received.")
    else:
        print(f"\n❌ Login failed. Run test_auth.py first to create user.")
        exit()


def auth_headers():
    return {"Authorization": f"Bearer {TOKEN}"}


# ─── CREATE TEST FILES ────────────────────────────────────

def create_test_pdf():
    try:
        from reportlab.pdfgen import canvas
        path = "test_lecture.pdf"
        c = canvas.Canvas(path)

        c.drawString(100, 750, "Introduction to Machine Learning")
        c.drawString(100, 700, "Machine learning is a subset of artificial intelligence.")
        c.drawString(100, 650, "It allows systems to learn from data automatically.")
        c.showPage()

        c.drawString(100, 750, "Types of Machine Learning")
        c.drawString(100, 700, "1. Supervised Learning")
        c.drawString(100, 650, "2. Unsupervised Learning")
        c.drawString(100, 600, "3. Reinforcement Learning")
        c.showPage()

        c.drawString(100, 750, "Applications of Machine Learning")
        c.drawString(100, 700, "- Image Recognition")
        c.drawString(100, 650, "- Natural Language Processing")
        c.drawString(100, 600, "- Recommendation Systems")
        c.showPage()

        c.save()
        return path
    except ImportError:
        print("   ⚠️  reportlab not installed. Run: pip install reportlab")
        return None


def create_test_docx():
    try:
        from docx import Document

        path = "test_lecture.docx"
        doc = Document()

        doc.add_heading("Introduction to Deep Learning", level=1)
        doc.add_paragraph("Deep learning is a subset of machine learning.")
        doc.add_paragraph("It uses neural networks with many layers.")

        doc.add_heading("Neural Networks", level=1)
        doc.add_paragraph("Neural networks are inspired by the human brain.")
        doc.add_paragraph("They consist of layers of interconnected nodes.")

        doc.add_heading("Training a Neural Network", level=1)
        doc.add_paragraph("Training involves forward pass and backpropagation.")
        doc.add_paragraph("The model adjusts weights to minimize loss.")

        doc.save(path)
        return path
    except ImportError:
        print("   ⚠️  python-docx not installed.")
        return None


# ─── TESTS ────────────────────────────────────────────────

def test_upload_docx():
    path = create_test_docx()
    if not path:
        return False

    with open(path, "rb") as f:
        response = requests.post(
            f"{BASE_URL}/upload/lecture",
            headers=auth_headers(),
            data={"title": "Deep Learning Lecture", "language": "en"},
            files={"file": (path, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )

    print_result("UPLOAD — DOCX lecture", response)
    if os.path.exists(path):
        os.remove(path)
    return response.status_code == 201


def test_upload_pdf():
    path = create_test_pdf()
    if not path:
        return False

    with open(path, "rb") as f:
        response = requests.post(
            f"{BASE_URL}/upload/lecture",
            headers=auth_headers(),
            data={"title": "ML Lecture PDF", "language": "en"},
            files={"file": (path, f, "application/pdf")}
        )

    print_result("UPLOAD — PDF lecture", response)
    if os.path.exists(path):
        os.remove(path)
    return response.status_code == 201


def test_upload_no_file():
    response = requests.post(
        f"{BASE_URL}/upload/lecture",
        headers=auth_headers(),
        data={"title": "Test", "language": "en"},
        files={"dummy": ("", b"", "text/plain")}  # force multipart, no real file
    )
    print_result("UPLOAD — no file (should fail)", response)
    return response.status_code == 400


def test_upload_no_title():
    response = requests.post(
        f"{BASE_URL}/upload/lecture",
        headers=auth_headers(),
        data={"language": "en"},
        files={"dummy": ("", b"", "text/plain")}  # force multipart, no title
    )
    print_result("UPLOAD — no title (should fail)", response)
    return response.status_code == 400


def test_upload_wrong_extension():
    response = requests.post(
        f"{BASE_URL}/upload/lecture",
        headers=auth_headers(),
        data={"title": "Test", "language": "en"},
        files={"file": ("test.txt", b"some text content", "text/plain")}
    )
    print_result("UPLOAD — wrong extension (should fail)", response)
    return response.status_code == 400


def test_upload_wrong_language():
    from docx import Document
    path = "test_lang.docx"
    doc = Document()
    doc.add_heading("Test", level=1)
    doc.add_paragraph("Test content")
    doc.save(path)

    with open(path, "rb") as f:
        response = requests.post(
            f"{BASE_URL}/upload/lecture",
            headers=auth_headers(),
            data={"title": "Test", "language": "fr"},  # invalid language
            files={"file": (path, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )

    print_result("UPLOAD — wrong language (should fail)", response)
    if os.path.exists(path):
        os.remove(path)
    return response.status_code == 400


def test_upload_no_token():
    response = requests.post(
        f"{BASE_URL}/upload/lecture",
        data={"title": "Test", "language": "en"},
        files={"dummy": ("", b"", "text/plain")}
    )
    print_result("UPLOAD — no token (should fail)", response)
    return response.status_code == 401


def test_get_lectures():
    response = requests.get(
        f"{BASE_URL}/upload/lectures",
        headers=auth_headers()
    )
    print_result("GET — all lectures", response)
    return response.status_code == 200


def test_get_lectures_no_token():
    response = requests.get(f"{BASE_URL}/upload/lectures")
    print_result("GET — lectures no token (should fail)", response)
    return response.status_code == 401


# ─── RUN ALL ──────────────────────────────────────────────

if __name__ == "__main__":
    print("\n" + "="*50)
    print("   StudyGPT — UPLOAD MODULE TESTS")
    print("="*50)

    login()

    results = {}
    results["Upload DOCX"]            = test_upload_docx()
    results["Upload PDF"]             = test_upload_pdf()
    results["Upload no file"]         = test_upload_no_file()
    results["Upload no title"]        = test_upload_no_title()
    results["Upload wrong extension"] = test_upload_wrong_extension()
    results["Upload wrong language"]  = test_upload_wrong_language()
    results["Upload no token"]        = test_upload_no_token()
    results["Get all lectures"]       = test_get_lectures()
    results["Get lectures no token"]  = test_get_lectures_no_token()

    print("\n" + "="*50)
    print("   RESULTS SUMMARY")
    print("="*50)
    passed = failed = 0
    for test, result in results.items():
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"   {status} — {test}")
        if result:
            passed += 1
        else:
            failed += 1

    print(f"\n   Total: {passed} passed, {failed} failed")
    print("="*50)