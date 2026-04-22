import requests
import json
import os

BASE_URL = "http://localhost:5000/api"
TOKEN       = ""
LECTURE_ID  = ""  # will be set after upload


def print_result(test_name, response):
    print(f"\n{'='*50}")
    print(f"🔍 {test_name}")
    print(f"   Status Code : {response.status_code}")
    try:
        data = response.json()
        text = json.dumps(data, indent=2, ensure_ascii=False)
        print(f"   Response    : {text[:300]}")
    except:
        print(f"   Response    : {response.text[:300]}")
    print(f"{'='*50}")


# ─── LOGIN ────────────────────────────────────────────────

def login():
    global TOKEN
    response = requests.post(f"{BASE_URL}/auth/login", json={
        "email": "testuser@studygpt.com",
        "password": "Test1234!"
    })
    if response.status_code == 200:
        TOKEN = response.json().get("token")
        print(f"\n✅ Logged in. Token received.")
    else:
        print(f"\n❌ Login failed.")
        exit()


def auth_headers():
    return {"Authorization": f"Bearer {TOKEN}"}


# ─── UPLOAD TEST LECTURE ──────────────────────────────────

def upload_test_lecture():
    global LECTURE_ID
    from docx import Document

    path = "test_study_lecture.docx"
    doc = Document()

    doc.add_heading("Introduction to Artificial Intelligence", level=1)
    doc.add_paragraph("Artificial Intelligence is the simulation of human intelligence by machines.")
    doc.add_paragraph("AI systems can perform tasks that require human intelligence such as visual perception.")
    doc.add_paragraph("The field was founded in 1956 at Dartmouth College.")

    doc.add_heading("Machine Learning", level=1)
    doc.add_paragraph("Machine learning is a subset of AI that allows systems to learn from data.")
    doc.add_paragraph("It uses algorithms to parse data, learn from it and make decisions.")
    doc.add_paragraph("Types include supervised, unsupervised and reinforcement learning.")

    doc.add_heading("Deep Learning", level=1)
    doc.add_paragraph("Deep learning uses neural networks with many layers.")
    doc.add_paragraph("It excels in image recognition, speech recognition and natural language processing.")
    doc.add_paragraph("Deep learning requires large amounts of data and computational power.")

    doc.save(path)

    with open(path, "rb") as f:
        response = requests.post(
            f"{BASE_URL}/upload/lecture",
            headers=auth_headers(),
            data={"title": "Introduction to AI", "language": "en"},
            files={"file": (path, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )

    if os.path.exists(path):
        os.remove(path)

    if response.status_code == 201:
        LECTURE_ID = response.json().get("lecture_id")
        print(f"\n✅ Lecture uploaded | lecture_id={LECTURE_ID}")
        return True
    else:
        print(f"\n❌ Upload failed: {response.json()}")
        return False


# ─── TESTS ────────────────────────────────────────────────

def test_summary():
    response = requests.get(
        f"{BASE_URL}/study/summary/{LECTURE_ID}",
        headers=auth_headers()
    )
    print_result("STUDY — get summary", response)
    return response.status_code == 200


def test_explain_slide():
    response = requests.get(
        f"{BASE_URL}/study/explain/{LECTURE_ID}/1",
        headers=auth_headers()
    )
    print_result("STUDY — explain slide 1", response)
    return response.status_code == 200


def test_explain_invalid_slide():
    response = requests.get(
        f"{BASE_URL}/study/explain/{LECTURE_ID}/999",
        headers=auth_headers()
    )
    print_result("STUDY — explain invalid slide (should fail)", response)
    return response.status_code == 404


def test_flashcards():
    response = requests.get(
        f"{BASE_URL}/study/flashcards/{LECTURE_ID}",
        headers=auth_headers()
    )
    print_result("STUDY — get flashcards", response)
    return response.status_code == 200


def test_mindmap():
    response = requests.get(
        f"{BASE_URL}/study/mindmap/{LECTURE_ID}",
        headers=auth_headers()
    )
    print_result("STUDY — get mind map", response)
    return response.status_code == 200


def test_glossary():
    response = requests.get(
        f"{BASE_URL}/study/glossary/{LECTURE_ID}",
        headers=auth_headers()
    )
    print_result("STUDY — get glossary", response)
    return response.status_code == 200


def test_tldr():
    response = requests.get(
        f"{BASE_URL}/study/tldr/{LECTURE_ID}/1",
        headers=auth_headers()
    )
    print_result("STUDY — get tldr slide 1", response)
    return response.status_code == 200


def test_no_token():
    response = requests.get(f"{BASE_URL}/study/summary/{LECTURE_ID}")
    print_result("STUDY — no token (should fail)", response)
    return response.status_code == 401


def test_wrong_lecture_id():
    response = requests.get(
        f"{BASE_URL}/study/summary/00000000-0000-0000-0000-000000000000",
        headers=auth_headers()
    )
    print_result("STUDY — wrong lecture id (should fail)", response)
    return response.status_code == 404


# ─── RUN ALL ──────────────────────────────────────────────

if __name__ == "__main__":
    print("\n" + "="*50)
    print("   StudyGPT — STUDY CONTENT MODULE TESTS")
    print("="*50)

    login()

    uploaded = upload_test_lecture()
    if not uploaded:
        print("❌ Cannot run tests without lecture. Exiting.")
        exit()

    import time
    print(f"\n⏳ Waiting 5 seconds for RAG pipeline...")
    time.sleep(5)

    results = {}
    results["Summary"]               = test_summary()
    results["Explain slide"]         = test_explain_slide()
    results["Explain invalid slide"] = test_explain_invalid_slide()
    results["Flashcards"]            = test_flashcards()
    results["Mind map"]              = test_mindmap()
    results["Glossary"]              = test_glossary()
    results["TL;DR"]                 = test_tldr()
    results["No token"]              = test_no_token()
    results["Wrong lecture id"]      = test_wrong_lecture_id()

    print("\n" + "="*50)
    print("   RESULTS SUMMARY")
    print("="*50)
    passed = failed = 0
    for test, result in results.items():
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"   {status} — {test}")
        if result: passed += 1
        else: failed += 1

    print(f"\n   Total: {passed} passed, {failed} failed")
    print("="*50)