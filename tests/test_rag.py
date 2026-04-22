import sys
import os
import requests
import json

# Add project root to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

BASE_URL = "http://localhost:5000/api"
TOKEN    = ""


def print_result(test_name, passed, details=""):
    status = "✅ PASS" if passed else "❌ FAIL"
    print(f"\n{'='*50}")
    print(f"🔍 {test_name}")
    print(f"   Status : {status}")
    if details:
        print(f"   Details: {details}")
    print(f"{'='*50}")


# ─── LOGIN ────────────────────────────────────────────────

def login():
    global TOKEN
    response = requests.post(f"{BASE_URL}/auth/login", json={
        "email": "testuser@studygpt.com",
        "password": "Test1234!"
    })
    if response.status_code == 200:
        TOKEN = response.json().get("token")
        print(f"\n✅ Logged in. Token received.")
    else:
        print(f"\n❌ Login failed. Run test_auth.py first.")
        exit()


def auth_headers():
    return {"Authorization": f"Bearer {TOKEN}"}


# ─── CREATE TEST DOCX ─────────────────────────────────────

def create_test_docx():
    from docx import Document
    path = "test_rag_lecture.docx"
    doc = Document()

    doc.add_heading("Introduction to Machine Learning", level=1)
    doc.add_paragraph("Machine learning is a subset of artificial intelligence.")
    doc.add_paragraph("It enables systems to learn from data without being explicitly programmed.")
    doc.add_paragraph("Common applications include image recognition and natural language processing.")

    doc.add_heading("Supervised Learning", level=1)
    doc.add_paragraph("Supervised learning uses labeled training data.")
    doc.add_paragraph("The model learns to map inputs to outputs.")
    doc.add_paragraph("Examples include linear regression and decision trees.")

    doc.add_heading("Neural Networks", level=1)
    doc.add_paragraph("Neural networks are inspired by the human brain.")
    doc.add_paragraph("They consist of layers of interconnected neurons.")
    doc.add_paragraph("Deep learning uses neural networks with many hidden layers.")

    doc.add_heading("Overfitting", level=1)
    doc.add_paragraph("Overfitting occurs when model learns training data too well.")
    doc.add_paragraph("The model performs poorly on new unseen data.")
    doc.add_paragraph("Techniques like dropout and regularization help prevent overfitting.")

    doc.save(path)
    return path


# ─── TEST 1: UPLOAD AND TRIGGER RAG ──────────────────────

def test_upload_triggers_rag():
    """Upload a lecture and verify RAG pipeline runs"""
    path = create_test_docx()

    with open(path, "rb") as f:
        response = requests.post(
            f"{BASE_URL}/upload/lecture",
            headers=auth_headers(),
            data={"title": "ML RAG Test Lecture", "language": "en"},
            files={"file": (path, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )

    if os.path.exists(path):
        os.remove(path)

    passed = response.status_code == 201
    data   = response.json() if response.status_code == 201 else {}

    print_result(
        "UPLOAD — triggers RAG pipeline",
        passed,
        f"lecture_id={data.get('lecture_id')} | slides={data.get('total_slides')}"
    )

    return data.get("lecture_id") if passed else None


# ─── TEST 2: VERIFY EMBEDDINGS SAVED IN DB ───────────────

def test_embeddings_saved(lecture_id: str):
    """Check embeddings table has records for this lecture"""
    from dotenv import load_dotenv
    load_dotenv()

    try:
        from app import create_app
        from app.models.embedding import Embedding

        app = create_app()
        with app.app_context():
            count = Embedding.query.filter_by(lecture_id=lecture_id).count()
            passed = count > 0
            print_result(
                "EMBEDDINGS — saved in pgvector",
                passed,
                f"Embeddings found in DB: {count}"
            )
            return passed, count
    except Exception as e:
        print_result("EMBEDDINGS — saved in pgvector", False, str(e))
        return False, 0


# ─── TEST 3: RETRIEVER SIMILARITY SEARCH ─────────────────

def test_retriever(lecture_id: str):
    """Test retriever finds relevant chunks for a query"""
    try:
        from dotenv import load_dotenv
        load_dotenv()

        from app import create_app
        from app.modules.rag.retriever import retrieve_relevant_chunks

        app = create_app()
        with app.app_context():

            # Test 1: relevant query
            chunks = retrieve_relevant_chunks(
                query="What is overfitting in machine learning?",
                lecture_id=lecture_id,
                top_k=3
            )
            passed = len(chunks) > 0
            print_result(
                "RETRIEVER — relevant query",
                passed,
                f"Chunks returned: {len(chunks)} | Top slide: {chunks[0]['slide_number'] if chunks else 'None'} | Similarity: {chunks[0]['similarity'] if chunks else 'N/A'}"
            )

            # Test 2: another relevant query
            chunks2 = retrieve_relevant_chunks(
                query="Explain neural networks",
                lecture_id=lecture_id,
                top_k=3
            )
            passed2 = len(chunks2) > 0
            print_result(
                "RETRIEVER — neural networks query",
                passed2,
                f"Chunks returned: {len(chunks2)} | Top slide: {chunks2[0]['slide_number'] if chunks2 else 'None'} | Similarity: {chunks2[0]['similarity'] if chunks2 else 'N/A'}"
            )

            # Test 3: show retrieved content
            if chunks:
                print(f"\n   📄 Retrieved content preview:")
                for c in chunks:
                    print(f"      Slide {c['slide_number']} (sim={c['similarity']}): {c['content'][:80]}...")

            return passed and passed2

    except Exception as e:
        print_result("RETRIEVER — similarity search", False, str(e))
        return False


# ─── TEST 4: WRONG LECTURE ID ─────────────────────────────

def test_retriever_wrong_lecture():
    try:
        from dotenv import load_dotenv
        load_dotenv()

        from app import create_app
        from app.modules.rag.retriever import retrieve_relevant_chunks

        app = create_app()
        with app.app_context():
            chunks = retrieve_relevant_chunks(
                query="What is machine learning?",
                lecture_id="00000000-0000-0000-0000-000000000000",
                top_k=3
            )
            passed = len(chunks) == 0
            print_result(
                "RETRIEVER — wrong lecture id (should return empty)",
                passed,
                f"Chunks returned: {len(chunks)}"
            )
            return passed
    except Exception as e:
        print_result("RETRIEVER — wrong lecture id", False, str(e))
        return False


# ─── TEST 5: RE-UPLOAD CLEANS OLD EMBEDDINGS ─────────────

def test_reupload_cleans_embeddings(lecture_id: str):
    """Verify old embeddings deleted when lecture re-processed"""
    try:
        from dotenv import load_dotenv
        load_dotenv()

        from app import create_app
        from app.models.embedding import Embedding
        from app.modules.rag.services import process_lecture_embeddings

        app = create_app()
        with app.app_context():
            before = Embedding.query.filter_by(lecture_id=lecture_id).count()
            result = process_lecture_embeddings(lecture_id)
            after  = Embedding.query.filter_by(lecture_id=lecture_id).count()

            passed = result["success"] and after > 0
            print_result(
                "RAG — re-process cleans old embeddings",
                passed,
                f"Before: {before} | After: {after} | Success: {result['success']}"
            )
            return passed
    except Exception as e:
        print_result("RAG — re-process cleans old embeddings", False, str(e))
        return False


# ─── RUN ALL ──────────────────────────────────────────────

if __name__ == "__main__":
    print("\n" + "="*50)
    print("   StudyGPT — RAG PIPELINE TESTS")
    print("="*50)

    login()

    results = {}

    # Step 1: Upload lecture and trigger RAG
    lecture_id = test_upload_triggers_rag()
    results["Upload triggers RAG"] = lecture_id is not None

    if lecture_id:
        import time
        print(f"\n⏳ Waiting 3 seconds for RAG pipeline to complete...")
        time.sleep(3)

        # Step 2: Verify embeddings saved
        emb_passed, emb_count = test_embeddings_saved(lecture_id)
        results["Embeddings saved in DB"] = emb_passed

        # Step 3: Test retriever
        results["Retriever similarity search"] = test_retriever(lecture_id)

        # Step 4: Wrong lecture id
        results["Retriever wrong lecture id"] = test_retriever_wrong_lecture()

        # Step 5: Re-upload cleans embeddings
        results["Re-process cleans embeddings"] = test_reupload_cleans_embeddings(lecture_id)

    else:
        print("\n❌ Upload failed. Skipping remaining tests.")

    print("\n" + "="*50)
    print("   RESULTS SUMMARY")
    print("="*50)
    passed = failed = 0
    for test, result in results.items():
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"   {status} — {test}")
        if result: passed += 1
        else: failed += 1

    print(f"\n   Total: {passed} passed, {failed} failed")
    print("="*50)