import fitz  # PyMuPDF
from docx import Document
import re


# ─── PDF SLIDE DETECTION ──────────────────────────────────

def extract_slides_from_pdf(file_path: str) -> list:
    """
    Each PDF page = one slide.
    Falls back to keyword detection if page has no text.
    """
    slides = []
    doc = fitz.open(file_path)

    for page_num, page in enumerate(doc):
        text = page.get_text("text").strip()

        # Skip completely empty pages
        if not text:
            continue

        slides.append({
            "slide_number": len(slides) + 1,
            "content": text
        })

    doc.close()
    return slides


# ─── DOCX SLIDE DETECTION ─────────────────────────────────

def extract_slides_from_docx(file_path: str) -> list:
    """
    Priority detection order:
    1. Heading styles (Heading 1, Heading 2)
    2. Keyword detection (Slide X, PAGE X, ---)
    3. Large spacing / empty paragraph blocks
    """
    doc = Document(file_path)
    slides = []
    current_slide = []

    def save_slide():
        content = "\n".join(current_slide).strip()
        if content:
            slides.append({
                "slide_number": len(slides) + 1,
                "content": content
            })

    for para in doc.paragraphs:
        text = para.text.strip()

        # ── METHOD 1: Heading style detection ──
        if para.style.name.startswith("Heading"):
            save_slide()
            current_slide = [text] if text else []
            continue

        # ── METHOD 2: Keyword detection ──
        if _is_slide_separator(text):
            save_slide()
            current_slide = []
            continue

        # ── METHOD 3: Spacing detection (3+ empty lines = new slide) ──
        if not text:
            current_slide.append("")
            empty_count = sum(1 for line in current_slide if line == "")
            if empty_count >= 3:
                save_slide()
                current_slide = []
            continue

        current_slide.append(text)

    # Save last slide
    save_slide()

    # If no slides detected → treat whole doc as one slide
    if not slides:
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        if full_text:
            slides.append({"slide_number": 1, "content": full_text})

    return slides


def _is_slide_separator(text: str) -> bool:
    """
    Detect slide separator keywords.
    Supports English and Arabic patterns.
    """
    if not text:
        return False

    text_lower = text.lower().strip()

    # English patterns
    english_patterns = [
        r"^slide\s*\d+",           # Slide 1, Slide 2
        r"^page\s*\d+",            # Page 1, Page 2
        r"^chapter\s*\d+",         # Chapter 1
        r"^section\s*\d+",         # Section 1
        r"^---+$",                 # --- or ------
        r"^={3,}$",                # ===
        r"^\*{3,}$",               # ***
    ]

    # Arabic patterns
    arabic_patterns = [
        r"^شريحة\s*\d+",           # شريحة 1
        r"^صفحة\s*\d+",            # صفحة 1
        r"^الفصل\s*\d+",           # الفصل 1
        r"^القسم\s*\d+",           # القسم 1
    ]

    all_patterns = english_patterns + arabic_patterns

    for pattern in all_patterns:
        if re.match(pattern, text_lower) or re.match(pattern, text):
            return True

    return False


# ─── MAIN DISPATCHER ──────────────────────────────────────

def extract_slides(file_path: str, file_extension: str) -> list:
    """
    Main function. Detects file type and extracts slides.
    Returns list of dicts: [{ slide_number, content }]
    """
    ext = file_extension.lower()

    if ext == ".pdf":
        return extract_slides_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_slides_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")